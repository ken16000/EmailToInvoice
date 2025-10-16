import streamlit as st
import pandas as pd
from google import genai
from google.genai import types
import json
from io import BytesIO 
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# --- UI部分は省略 ---

st.title("✉️ メール本文から見積書を作成するアプリ")
st.markdown("お客様からの**見積もり依頼メール**と会社情報を入力し、Gemini APIで見積書データを作成します。")

# Google Gemini API Key
api_key = st.text_input("🔑 Gemini API Keyを入力", type="password")

# メール本文入力
email_body = st.text_area(
    "✉️ お客様からの見積もり依頼メール本文を貼り付け",
    height=300,
    placeholder="例:\n件名：〇〇プロジェクトの見積もり依頼\n株式会社XX\n田中様\n\nいつもお世話になっております。\nご説明いただいた〇〇システムの導入につき、以下の内容で見積もりをお願いいたします。\n\n・コンサルティング費用: 1式\n・システムライセンス: 50ライセンス\n\n納期は来月末、本見積もりの有効期限は発行日から1ヶ月を希望します。\n"
)

# 会社情報入力
company_info = st.text_area(
    "🏢 見積書発行元の会社情報を入力",
    height=150,
    placeholder="例:\n【見積元情報】\n会社名: △△合同会社\n住所: 東京都千代田区1-2-3\n電話: 03-1234-5678\n担当者名: 山田 太郎\n適格請求書発行事業者登録番号: T1234567890123"
)

# 見積書作成ボタン
if st.button("🚀 見積書データを生成"):
    if not api_key:
        st.error("Gemini API Keyを入力してください。")
    elif not email_body or not company_info:
        st.error("メール本文と会社情報の両方を入力してください。")
    else:
        try:
            # --- 2. Gemini APIの呼び出し (プロンプトとスキーマを更新) ---
            client = genai.Client(api_key=api_key)
            
            # プロンプトの定義 (JSONスキーマは変更なし)
            prompt = f"""
            以下のメール本文と見積書発行元の会社情報に基づき、見積書作成に必要なデータをJSON形式で抽出・生成してください。
            特に、明細の単価と金額は、一般的な市場価格や既知の価格に基づいて生成してください。
            返答は**JSONデータのみ**にしてください。
            
            【メール本文】
            {email_body}
            
            【見積書発行元の会社情報】
            {company_info}
            
            【JSONスキーマ】
            {{
              "発行日": "YYYY年MM月DD日",
              "見積書番号": "任意で生成",
              "見積先名": "メール本文から抽出",
              "見積先住所": "メール本文から可能な限り抽出",
              "見積元情報": "入力された会社情報をそのまま使用",
              "有効期限": "YYYY年MM月DD日",
              "納期": "〇〇日以内 または YYYY年MM月DD日",
              "明細": [
                {{"品目": "コンサルティング費用", "単価": 250000, "数量": 1, "単位": "式", "税区分": "税別"}}
              ],
              "合計金額_税抜": 500000,
              "合計金額_税込": 550000
            }}
            """
            
            # API呼び出し
            with st.spinner("Geminiが見積書データを生成中です..."):
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )

            # JSONデータのパース
            json_text = response.text.strip().lstrip('```json').rstrip('```').strip()
            invoice_data = json.loads(json_text) 
            
            st.success("✅ 見積書データが正常に生成されました！")
            st.json(invoice_data)
            
            # ----------------------------------------------------
            # --- 3. 見積書のフォーマット (Word文書の生成を更新) ---
            # ----------------------------------------------------

            def create_word_quotation(data):
                """JSONデータからWord文書を作成する"""
                document = Document()
                
                # --- 文書全体のスタイル設定 ---
                style = document.styles['Normal']
                style.font.name = '游ゴシック' 
                style.font.size = Pt(10.5)

                # --- タイトル ---
                title = document.add_heading('御 見 積 書', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.size = Pt(20)

                # --- 発行情報（右寄せ）---
                document.add_paragraph()
                p_date = document.add_paragraph(f'発行日： {data.get("発行日", "日付不明")}')
                p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # --- 請求先情報 ---
                document.add_paragraph()
                p_client = document.add_paragraph()
                p_client.add_run(f'{data.get("見積先名", "（見積先名不明）")}').bold = True
                p_client.add_run(' 様')
                document.add_paragraph(f'（ご担当者様名：【ご担当者名】）')
                document.add_paragraph()
                
                # --- 見積元情報 ---
                p_sender = document.add_paragraph()
                p_sender.add_run(f'【発行元】').bold = True
                p_sender.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # ★★★ エラー修正箇所 ★★★
                # 値が辞書型の場合は、改行で結合して文字列に変換する
                sender_info_raw = data.get("見積元情報", "")
                if isinstance(sender_info_raw, dict):
                    # 辞書の場合は、key: valueの形式で結合
                    sender_info = [f"{k}: {v}" for k, v in sender_info_raw.items()]
                else:
                    # 文字列の場合は、改行で分割
                    sender_info = str(sender_info_raw).split('\n')
                
                p_sender_detail = document.add_paragraph()
                for line in sender_info:
                     p_sender_detail.add_run(line + '\n')
                p_sender_detail.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                document.add_paragraph()


                # --- 合計金額 ---
                document.add_paragraph()
                p_total = document.add_paragraph()
                p_total.add_run('下記の通りお見積り申し上げます。').bold = True
                p_total.add_run(f'\n\nお見積金額合計: ¥{data.get("合計金額_税込", 0):,} (税込)').bold = True
                p_total.runs[1].font.size = Pt(16)
                document.add_paragraph()
                
                # --- 納期・有効期限 ---
                document.add_paragraph(f'■ 有効期限: {data.get("有効期限", "発行日より1ヶ月")}')
                document.add_paragraph(f'■ 納　　期: {data.get("納期", "別途協議")}')

                document.add_paragraph('\n')
                document.add_heading('■ 見積明細', 3)

                # --- 明細テーブル ---
                details = data.get("明細", [])
                if details:
                    table = document.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    hdr_cells = table.rows[0].cells
                    headers = ["品目", "単価", "数量", "単位", "金額 (税抜)"]
                    for i, header in enumerate(headers):
                        cell = hdr_cells[i]
                        cell.text = header
                        cell.paragraphs[0].runs[0].font.bold = True
                        
                    # データ行
                    for item in details:
                        row_cells = table.add_row().cells
                        row_cells[0].text = item.get("品目", "")
                        row_cells[1].text = f'¥{item.get("単価", 0):,}'
                        row_cells[2].text = str(item.get("数量", ""))
                        row_cells[3].text = item.get("単位", "")
                        
                        # 金額計算 (単価 * 数量)
                        amount = item.get("単価", 0) * item.get("数量", 0)
                        row_cells[4].text = f'¥{amount:,}'
                
                # 合計情報（明細テーブルの下に追記）
                document.add_paragraph()
                document.add_paragraph(f'【小計 (税抜)】: ¥{data.get("合計金額_税抜", 0):,}')
                # 税額の計算
                tax_amount = data.get("合計金額_税込", 0) - data.get("合計金額_税抜", 0)
                document.add_paragraph(f'【消費税 (10%)】: ¥{tax_amount:,}')
                document.add_paragraph(f'【合計金額 (税込)】: ¥{data.get("合計金額_税込", 0):,}').bold = True


                # 文書をBytesIOに保存
                output = BytesIO()
                document.save(output)
                return output.getvalue()

            word_data = create_word_quotation(invoice_data) # 変数名はそのまま利用

            # --- 4. ダウンロード機能の実装 (Word形式で提供) ---
            st.markdown("### 📥 見積書データ（Word形式）ダウンロード")
            
            st.download_button(
                label="見積書をダウンロード (Word .docx)",
                data=word_data,
                file_name=f'quotation_{pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                help="Geminiが抽出したデータに基づき、Word文書を生成します。"
            )

            st.info("💡 ダウンロードしたWordファイルを開き、レイアウトを調整してご利用ください。")

        except json.JSONDecodeError:
            st.error("❌ Geminiからの応答がJSON形式ではありませんでした。プロンプトを見直すか、再度お試しください。")
            st.code(response.text)
        except Exception as e:
            st.error(f"❌ エラーが発生しました: {e}")
            st.info("APIキーやデータ形式が正しいか、ネットワーク接続を確認してください。")